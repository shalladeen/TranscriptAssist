import streamlit as st
import pandas as pd
from docx import Document
import spacy
import re
import io
import json
import datetime

nlp = spacy.load("en_core_web_sm")
df = pd.DataFrame()

# Function to extract text from a Word document
def extract_speaker_blocks(docx_file):
    doc = Document(docx_file)
    speaker_pattern = re.compile(r"^([A-Z][a-z]+,\s[A-Z][a-z]+(?:\s[A-Z]\.)?)\s+(\d{1,2}:\d{2}(?::\d{2})?)")

    blocks = []
    current_speaker = None

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue

        match = speaker_pattern.match(line)
        if match:
            current_speaker = match.group(1)
        elif current_speaker:
            blocks.append((current_speaker, line))

    return blocks

# Spacy Function
def is_action_item_spacy(line):
    doc = nlp(line)

    # Look for modal verbs or future tense (will, should, can, etc.)
    for token in doc:
        if token.tag_ in ("MD", "VBP", "VB") and token.dep_ == "ROOT":
            verb = token.lemma_.lower()
            if verb in ["share", "talk", "assign", "send", "remind", "follow", "check", "email", "connect", "complete", "submit", "update"]:
                return True
    return False

# Rejex trigger patterns
ACTION_PATTERNS = [
    r"\b(can you|please|make sure|follow up|remind|send|complete|schedule|take care of|ensure|need to|have to|required to|I'll|I will|let's|should)\b.*",
]

# Regex for name detection (Lastname, Firstname format)
NAME_PATTERN = re.compile(r"\b[A-Z][a-z]+,\s[A-Z][a-z]+(?:\s[A-Z]\.)?\b")

# Function to find action items in transcript
def find_action_items_with_speakers(speaker_blocks, keywords):
    action_items = []

    for speaker, line in speaker_blocks:
        line_clean = line.strip()
        if not line_clean or len(line_clean.split()) < 4:
            continue  # skip short/noisy lines

        # Detection logic
        keyword_match = any(keyword.lower() in line_clean.lower() for keyword in keywords)
        regex_match = any(re.search(pattern, line_clean, re.IGNORECASE) for pattern in ACTION_PATTERNS)
        spacy_match = is_action_item_spacy(line_clean)

        if keyword_match or regex_match or spacy_match:
            item_type = "Confirmed"
        else:
            # Fallback: if line has a verb or a soft signal
            doc = nlp(line_clean)
            has_verb = any(tok.pos_ == "VERB" for tok in doc)
            has_soft_cue = re.search(r"\b(I|We|Let's|Can|Should|Maybe|Please|Need)\b", line_clean, re.IGNORECASE)

            if has_verb or has_soft_cue:
                item_type = "Possible"
            else:
                continue  # skip irrelevant

        action_items.append({
            "Action Item": line_clean,
            "Type": item_type
        })

    return pd.DataFrame(action_items) if action_items else pd.DataFrame(columns=["Action Item", "Type"])


# Extract full and first names from speaker tags
def extract_names(docx_file):
    doc = Document(docx_file)
    name_pattern = re.compile(r"\b([A-Z][a-z]+,\s[A-Z][a-z]+(?:\s[A-Z]\.)?)\b")

    full_names = set()
    first_names = set()

    for para in doc.paragraphs:
        line = para.text.strip()
        match = name_pattern.search(line)
        if match:
            full_name = match.group(1)
            full_names.add(full_name)
            first = full_name.split(",")[1].strip().split(" ")[0]
            first_names.add(first)

    return sorted(full_names), sorted(first_names)

        
# Streamlit UI
st.set_page_config(page_title="Transcript Action Item Extractor", layout="wide")

# Drag and Drop Style
st.markdown("""
    <style>
    .stFileUploader > div {
        border: 2px dashed #CBC3E6 !important;
        border-radius: 8px;
        background-color: #F2F1F9;
        transition: background-color 0.3s ease;
    }
    .stFileUploader > div:hover {
        background-color: #E6E3F3 !important;
    }
    .stTextInput>div>div>input,
    .stTextArea textarea {
    resize: none !important;
    background-color: white;
    border: 1px solid #E8E8E8;
    border-radius: 6px;
    padding: 0.5rem;
    font-size: 0.95rem;
}

    .stButton>button, .stDownloadButton>button {
        background-color: #5236AB;
        color: white;
        border-radius: 5px;
        padding: 8px 16px;
        font-weight: 500;
        border: none;
    }
    .stButton>button:hover, .stDownloadButton>button:hover {
        background-color: #9E83F5;
    }
    </style>
""", unsafe_allow_html=True)

# Header & intro with bottom border
st.markdown("""
    <div style='border-bottom: 1px solid #E8E8E8; padding-bottom: 0.5rem; margin-bottom: 1rem;'>
        <h2 style='margin-bottom: 0.2rem;'>Extract Action Items From Your Transcript</h2>
        <p style='color: #666666; font-size: 0.9rem;'>
            Upload your <code>.docx</code> meeting transcript and extract actionable tasks with speaker attribution.
            Export results in CSV, JSON, or Markdown.
        </p>
    </div>
""", unsafe_allow_html=True)

# Upload
st.markdown("### Upload Your File")
uploaded_file = st.file_uploader("Choose a Word document (.docx)", type=["docx"])

# Detection tuning section
st.markdown("### Refine Detection")
keywords = st.text_area(
    "Comma-separated keywords",
    "action, follow up, send, complete, email, share, remind, assign, connect, confirm"
).split(",")

# Main logic
if uploaded_file:
    with st.spinner("Analyzing transcript..."):
        speaker_blocks = extract_speaker_blocks(uploaded_file)
        _, first_names = extract_names(uploaded_file)
        df = find_action_items_with_speakers(speaker_blocks, keywords)

    if not df.empty:
        st.success(f"Found {len(df)} action item(s).")
        st.markdown("### Extracted Action Items")
        st.dataframe(df, height=300, use_container_width=True)

        #  # Scroll hint - Just uncomment if wanted
        # st.markdown("""
        # <div style="text-align:center; margin-top: 0.5rem; color: #666666; font-size: 0.9rem;">
        #     Scroll down to download your results ↓
        # </div>
        # """, unsafe_allow_html=True)

        st.markdown("### Export Results")

        # Download button row
        col1, col2, col3 = st.columns([1, 1, 1])

        with col1:
            csv_data = df.to_csv(index=False).encode('utf-8')
            st.download_button("Download CSV", csv_data, "action_items.csv", "text/csv")

        with col2:
            json_data = df.to_json(orient='records', indent=4)
            st.download_button("Download JSON", json_data, "action_items.json", "application/json")

        with col3:
            md_output = "\n".join([f"- **{row['Action Item']}** _(Type: {row['Type']})_" for _, row in df.iterrows()])
            st.download_button("Download Markdown", md_output, "action_items.md", "text/markdown")
    else:
        st.info("No action items detected in the transcript.")

