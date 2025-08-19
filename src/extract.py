from typing import List, Tuple
from docx import Document
from .constants import SPEAKER_LINE_PATTERN, NAME_CAPTURE_PATTERN

def extract_speaker_blocks(docx_file) -> List[Tuple[str, str]]:
    doc = Document(docx_file)
    blocks = []
    current_speaker = None

    for para in doc.paragraphs:
        line = (para.text or "").strip()
        if not line:
            continue

        match = SPEAKER_LINE_PATTERN.match(line)
        if match:
            current_speaker = match.group(1)
        elif current_speaker:
            blocks.append((current_speaker, line))

    return blocks

def extract_names(docx_file):
    doc = Document(docx_file)

    full_names = set()
    first_names = set()

    for para in doc.paragraphs:
        line = (para.text or "").strip()
        match = NAME_CAPTURE_PATTERN.search(line)
        if match:
            full_name = match.group(1)
            full_names.add(full_name)
            first = full_name.split(",")[1].strip().split(" ")[0]
            first_names.add(first)

    return sorted(full_names), sorted(first_names)
